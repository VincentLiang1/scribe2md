r"""Office 新格式(docx / pptx / xlsx / xlsm)→ Block 清單。

三個函式庫都是**惰性載入**(模組層 None 佔位 + `_ensure_*()` 回填 global,
同 transcribe._ensure_whisper):啟動路徑不付 XML 解析器的 import 成本,
測試也才 monkeypatch 得掉。

失真一律標註,不安靜丟棄——合併儲存格、底色語意、圖表、追蹤修訂等,
在 md 裡都會留下〔〕標記並進 frontmatter 的 lossy_kinds。理由見 docmd。
"""
import bisect
import hashlib
import html
import logging
import re
import zipfile
from pathlib import Path

from meeting_scribe import docimage, docmd
from meeting_scribe.docmd import AssetsDir, Block, Heading, Image, Note, Para, Raw, Records, Table
from meeting_scribe.errors import UserFacingError

logger = logging.getLogger(__name__)

# 惰性載入的佔位(測試 monkeypatch 這三個屬性換假貨)
docx = None
pptx = None
openpyxl = None

# Office Open XML 的命名空間。python-docx/pptx 沒有把「依文件順序走」
# 與「取內嵌圖」包成公開 API,只能自己走 lxml 元素
_NS_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_NS_A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
_NS_R = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
# PresentationML(母片的 p:txStyles 定義了內文佔位符預設有沒有項目符號)
_NS_P = "{http://schemas.openxmlformats.org/presentationml/2006/main}"
# 舊版 Word 的 VML(內嵌 OLE 物件的預覽圖也走這條)
_NS_V = "{urn:schemas-microsoft-com:vml}"
# SmartArt:內容在另一個 part,主文件只留 <dgm:relIds> 指標
_NS_DGM = "{http://schemas.openxmlformats.org/drawingml/2006/diagram}"

# 工作表大小的降級門檻(格數)。超過就放棄合併/顏色/圖表的逐格檢查,只串流
# 讀值——8GB 基準機上 openpyxl 的正常模式讀大表可以吃到 GB 級記憶體。
# **資料一列都不截斷**,降級的只有 metadata 檢查
_HUGE_CELLS = 2_000_000

# 圖表型別的繁中說法(給人看的 caption;認不出來就不寫,不要硬翻)
_CHART_KIND_ZH = {
    "PIE": "圓餅圖", "DOUGHNUT": "環圈圖", "BAR": "長條圖",
    "COLUMN": "直條圖", "LINE": "折線圖", "AREA": "面積圖",
    "XY_SCATTER": "散佈圖", "RADAR": "雷達圖",
}

# 單一合併範圍的展開上限:`A1:XFD1048576` 這種整表合併若照實展開會產生
# 上百億個 key 直接把記憶體吃光(真實檔案中確實存在誤操作的巨大合併)
_MAX_MERGE_CELLS = 100_000

# 內容控制項可以巢狀(表單裡包表單)。上限只是防呆
_MAX_SDT_DEPTH = 8


def _ensure_docx():
    global docx
    if docx is None:
        docx = docmd.lazy_import("docx", "Word 檔")
    return docx


def _ensure_pptx():
    global pptx
    if pptx is None:
        pptx = docmd.lazy_import("pptx", "PowerPoint 檔")
    return pptx


def _ensure_openpyxl():
    global openpyxl
    if openpyxl is None:
        openpyxl = docmd.lazy_import("openpyxl", "Excel 檔")
    return openpyxl


def _clean(text) -> str:
    r"""正規化一段文字。

    **`\v` 一定要換成 `\n`**:PowerPoint 的軟換行(`<a:br/>`)被 python-pptx
    給成垂直定位字元 `\x0b`,原樣寫進 md 就是一個看不見的控制字元
    (2026-08-02 實測踩到:一段內文在編輯器裡看起來斷了行,實際上是 \x0b)。
    Word 不會產生它,所以放在共用的清理函式裡沒有副作用。"""
    text = re.sub(r"[\v\f]", "\n", str(text or ""))
    return re.sub(r"[ \t]+\n", "\n", text).strip()


# ---- 超連結(docx / pptx 共用)----

def _external_links(part) -> dict[str, str]:
    """rId → 外部網址,**只收 `TargetMode="External"` 的**。

    每個 part 各有自己的一份 rels(內文、每個頁首頁尾、每張投影片、每張
    備忘稿),所以查表一定要拿**那個 part 自己的**——共用一份的話,同一個
    rId 在不同 part 指向不同東西,連結會接錯。

    內部的 rels(圖片、註腳、佈景主題)照樣在這張表外面:只有掛在
    `w:hyperlink` / `a:hlinkClick` 上的 rId 才會來查,不會誤取。"""
    rels = getattr(part, "rels", None)
    if not rels:
        return {}
    return {
        rid: str(rel.target_ref)
        for rid, rel in rels.items()
        if getattr(rel, "is_external", False) and getattr(rel, "target_ref", "")
    }


def _as_link(text: str, url: str) -> str:
    r"""文字 + 網址 → markdown 連結(2026-08-04 使用者指定直接寫成連結)。

    在此之前,連結的**文字**進得了 md(`_para_text` 收所有 `w:t`)、**網址**
    整個丟掉,而且不留標記——實測 424 份 docx 有 83 份(20%)、194 份 pptx
    有 46 份(24%)帶外部連結,最多的一份 90 個。那種文件的網址往往就是
    它的參考文獻,無聲丟掉是實打實的內容損失。

    三個細節:
    - **前後空白留在連結外面**:`[ 官網 ](…)` 雖然也能渲染,但多出來的
      空白會跑進連結文字裡。
    - **文字與網址相同就寫成自動連結** `<https://…>`:`[https://a](https://a)`
      是同一個字串寫兩次,對 RAG 只是多一份雜訊。
    - **網址含空白或括號要包進 `<>`**:那是 CommonMark 的寫法,不包的話
      連結會在第一個空白處斷掉(Word 允許使用者貼進帶空白的網址)。"""
    core = text.strip()
    url = (url or "").strip()
    if not core or not url:
        return text
    lead = text[:len(text) - len(text.lstrip())]
    tail = text[len(text.rstrip()):]
    if core == url:
        return f"{lead}<{url}>{tail}"
    if re.search(r"[\s()<>]", url):
        url = "<" + url.replace("<", "%3C").replace(">", "%3E") + ">"
    # 文字裡的方括號會提前結束連結。反斜線跳脫是 markdown 的標準寫法,
    # 而稽核比對那側本來就會把反斜線壓掉(docmd.squash)
    core = core.replace("[", r"\[").replace("]", r"\]")
    return f"{lead}[{core}]({url}){tail}"


# ---- docx ----

def _heading_level(para) -> int:
    """段落樣式 → 標題階層(0 = 不是標題)。

    中文版 Word 的內建樣式名是「標題 1」而不是 "Heading 1",兩種都要認
    ——只認英文的話,同仁用中文版 Word 寫的文件會整份變成無階層的平坦
    段落,RAG 切塊就沒有依據了。"""
    style = getattr(para, "style", None)
    name = str(getattr(style, "name", "") or "")
    m = re.match(r"^\s*(?:Heading|標題)\s*(\d+)", name)
    if m:
        return min(int(m.group(1)), 6)
    if name.strip() in ("Title", "標題"):
        return 1
    return 0


def _image_rids(element) -> list[str]:
    r"""一段 XML 裡所有內嵌圖的 rId,依出現順序、去重。

    **兩種標記都要認**:`<a:blip r:embed>` 是 DrawingML(Word 2007 之後),
    `<v:imagedata r:id>` 是舊的 VML——舊版 Word 存的圖、以及**內嵌 OLE 物件
    (Visio 圖、Excel 圖表)的預覽圖**都走 VML。只認前者的話,2018 年的公司
    文件常常一張圖都抽不到(使用者 2026-08-02 回報的那份:5 張圖裡 4 張是
    VML/OLE,整份 md 沒有任何圖片)。"""
    rids: list[str] = []
    for tag, attr in (
        (f"{_NS_A}blip", f"{_NS_R}embed"),
        (f"{_NS_V}imagedata", f"{_NS_R}id"),
    ):
        for node in element.iter(tag):
            rid = node.get(attr)
            if rid and rid not in rids:
                rids.append(rid)
    return rids


def _para_text(element, links: dict[str, str] | None = None) -> str:
    r"""段落文字,**含追蹤修訂的插入**(`w:ins`)、超連結網址與註腳記號。

    `python-docx` 的 `paragraph.text` 只取 `w:p` **直屬**的 `w:r/w:t`,而
    修訂插入的內容包在 `<w:ins><w:r><w:t>` 底下、深一層,整段被跳過。
    實測一份委外合約(使用者 2026-08-02 的全碟稽核揪出來):

        para.text  自民國11年1月1日起至民國11年12月31日止
        實際內容    自民國115年1月1日起至民國115年12月31日止

    **數字被無聲吃掉一位,而且讀起來完全合理**(民國 11 年 = 1922 年)
    ——這比整段消失更危險,因為沒有任何跡象顯示它錯了。那份檔案 27 段
    受影響,而我們還印著「以下內容是接受所有修訂後的最終版」,那句話在
    修好之前是假的。

    **刪除的文字用 `w:delText` 而不是 `w:t`**,所以「收集所有 w:t」正好
    等於「接受所有修訂」,不必另外判斷。文字方塊(`w:txbxContent`)則跳過
    ——那條由 `_textbox_texts` 另外處理並自成一段,混進本文只會讓兩段話
    交錯在一起。

    `links`(rId → 外部網址,由 `_external_links` 對**這個 part** 算好)給了
    才會把 `w:hyperlink` 包住的那幾個 run 併成一個 markdown 連結。**沒有
    `r:id` 的不算**:目錄與交互參照用的是 `w:anchor`(文件內錨點),實測
    424 份 docx 裡這種有 4,716 個、真正的外部連結才 858 個,把錨點也寫成
    連結只會產出一堆指不到任何地方的 `[文字]()`。

    註腳/章節附註的參照(`w:footnoteReference`)在這裡換成 `[註3]` 記號,
    正文與 `_note_blocks` 收在文末的內容靠這個編號對得起來。"""
    parts: list[str] = []
    pending: list[str] = []  # 目前這個 w:hyperlink 底下累積的文字
    pending_url = ""
    current = None  # 目前所在的 w:hyperlink 元素(None = 不在連結裡)

    def flush() -> None:
        nonlocal pending, pending_url
        if pending:
            parts.append(_as_link("".join(pending), pending_url))
        pending = []
        pending_url = ""

    for node in element.iter(
        f"{_NS_W}t", f"{_NS_W}footnoteReference", f"{_NS_W}endnoteReference",
    ):
        # 一次走完祖先鏈:同時要知道「在不在文字方塊裡」與「在不在連結裡」
        hyperlink = None
        boxed = False
        for anc in node.iterancestors():
            if anc.tag == f"{_NS_W}txbxContent":
                boxed = True
                break
            if hyperlink is None and anc.tag == f"{_NS_W}hyperlink":
                hyperlink = anc
        if boxed:
            continue
        # `iter` 是文件順序,所以「同一個 w:hyperlink 祖先」的連續節點就是
        # 同一個連結的文字(一個連結常被拆成好幾個 run:換字型、拼字檢查)
        if hyperlink is not current:
            flush()
            current = hyperlink
            rid = hyperlink.get(f"{_NS_R}id") if hyperlink is not None else None
            pending_url = (links or {}).get(rid or "", "")
        if node.tag == f"{_NS_W}t":
            text = node.text or ""
        else:
            mark = (
                docmd.FOOTNOTE_MARK if node.tag == f"{_NS_W}footnoteReference"
                else docmd.ENDNOTE_MARK
            )
            wid = node.get(f"{_NS_W}id")
            text = docmd.note_marker(mark, wid) if wid else ""
        (pending if hyperlink is not None else parts).append(text)
    flush()
    return "".join(parts)


def _related_blob(part, rid: str) -> bytes | None:
    """依 rId 取出關聯 part 的位元組。

    **python-docx 與 python-pptx 的 API 不同名**:前者是 `related_parts`
    (dict),後者是 `related_part(rId)`(方法),舊版還只有
    `rels[rId].target_part`。只寫其中一種的話,另一邊會安靜地什麼都抓不到
    ——SmartArt 那條就是這樣先在 pptx 上啞掉(實測 0/6)。"""
    for get in (
        lambda: part.related_parts[rid],
        lambda: part.related_part(rid),
        lambda: part.rels[rid].target_part,
    ):
        try:
            return get().blob
        except (KeyError, AttributeError, TypeError):
            continue
    logger.debug("取不到關聯 part rId=%s", rid)
    return None


def _diagram_texts(element, part) -> list[str]:
    r"""SmartArt(`dgm:relIds`)裡的文字。

    SmartArt 的內容不在 document.xml / slideN.xml 裡,而在另一個 part
    (`diagrams/dataN.xml`),主文件只留一個 `<dgm:relIds r:dm="rIdN">` 指標
    ——所以 python-docx/pptx 的一般走訪完全看不到它。實際案例(2026-08-02
    的全碟稽核:360 個檔中招):一張「跨業整合架構」流程圖的
    「第一步 完成集團跨業合作申請,建立法遵同意鎖控管機制」整組消失,
    而那正是那一頁的內容。

    `r:dm` 指的是**資料** part(另外還有 layout/style/colors 三個,那些是
    外觀不是內容)。文字節點與一般 DrawingML 一樣是 `<a:t>`。"""
    out: list[str] = []
    for node in element.iter(f"{_NS_DGM}relIds"):
        rid = node.get(f"{_NS_R}dm")
        if not rid:
            continue
        blob = _related_blob(part, rid)
        if blob is None:
            continue
        xml = blob.decode("utf-8", "replace")
        seen: list[str] = []
        for raw in re.findall(r"<a:t>([^<]*)</a:t>", xml):
            text = html.unescape(raw).strip()
            # 同一段文字在 data part 裡常常重複(版面各處引用同一個節點),
            # 去重但保序——重複三次的條列讀起來像壞掉
            if text and text not in seen:
                seen.append(text)
        if seen:
            out.append(" / ".join(seen))
    return out


def _textbox_texts(element) -> list[str]:
    """一段 XML 裡文字方塊(`w:txbxContent`)的文字,一個方塊一段。

    **python-docx 的 `paragraph.text` 看不到它們**:那只取 `w:p` 直屬的
    `w:r/w:t`,而文字方塊的內容是巢狀在
    `w:pict/v:shape/v:textbox/w:txbxContent` 底下的**另一組 `w:p`**。
    簡報式的 Word 文件很愛把圖說與架構圖的標籤全放在文字方塊裡——實際
    案例(使用者 2026-08-02 的稽核):「圖表 2:TWDT-ETH技術架構圖」連同
    一張架構圖的所有標籤(以太坊區塊鏈技術、去中心化交易、P2P網路…)
    整組消失,而那是圖的唯一文字說明。"""
    out: list[str] = []
    for box in element.iter(f"{_NS_W}txbxContent"):
        parts = [node.text or "" for node in box.iter(f"{_NS_W}t")]
        text = _clean("".join(parts))
        if text:
            out.append(text)
    return out


def _para_images(
    para, document, assets: AssetsDir, ocr_enabled: bool = True,
) -> list[Block]:
    """段落裡的內嵌圖 → 寫進 assets、送 OCR,並保住圖片在文件中的位置。"""
    return _images_from(para._p, document, assets, ocr_enabled)


def _images_from(
    element, document, assets: AssetsDir, ocr_enabled: bool = True,
) -> list[Block]:
    """任一段 XML 裡的內嵌圖 → Blocks(段落與表格共用)。"""
    out: list[Block] = []
    for rid in _image_rids(element):
        try:
            part = document.part.related_parts[rid]
            data = part.blob
            suffix = Path(str(getattr(part, "partname", "img.png"))).suffix or ".png"
        except (KeyError, AttributeError):
            logger.debug("docx 內嵌圖取不到 rId=%s,略過", rid)
            continue
        # emf/wmf 直接存的話 md 連結指向一個畫不出來的檔案(見 to_displayable)
        data, suffix = docimage.to_displayable(data, suffix)
        link = assets.add_bytes(data, suffix)
        if link:
            out.append(Image(link, "文件內嵌圖片"))
        # 圖裡的文字才是給 AI 用的重點:只留連結等於什麼都沒給
        out.extend(docimage.ocr_image_bytes(data, "文件裡的內嵌圖片", ocr_enabled))
    return out


def _docx_table_rows(table) -> tuple[list[list[str]], bool]:
    """docx 表格 → 純文字列。回傳 (列, 是否含巢狀表格)。

    markdown 表格不能巢狀,內層表只能攤平成同一格裡的多行文字。

    連結表從 `table.part` 自己取(內文與各個頁首頁尾是不同的 part,而
    表格是頁首排版最常用的東西)——用表格排版的文件很常見,格子裡的
    連結不能因為「它在表格裡」就掉一半。"""
    rows: list[list[str]] = []
    nested = False
    links = _external_links(getattr(table, "part", None))
    for row in table.rows:
        cells: list[str] = []
        for cell in row.cells:
            # 同 _para_text:cell.text 一樣漏掉 w:ins 裡的修訂插入
            text = _clean("\n".join(_para_text(x._p, links) for x in cell.paragraphs))
            if getattr(cell, "tables", None):
                nested = True
                for inner in cell.tables:
                    inner_rows, _ = _docx_table_rows(inner)
                    flat = " / ".join(" ".join(r) for r in inner_rows)
                    text = f"{text}\n{flat}".strip()
            cells.append(text)
        rows.append(cells)
    return rows, nested


# 自我稽核:一段文字要多長才值得比對。太短的片段(編號、單一字母、
# 表格裡的「-」)在輸出裡本來就會被拆散或合併,比對只會製造雜訊
def _xml_paragraphs(src: Path, part_pattern: str, para_tag: str, text_tag: str) -> list[str]:
    """直接從 zip 讀出「原始檔宣稱有哪些段文字」。

    **刻意不經 python-docx/pptx**:它們看不到的節點正是這裡要抓的東西。"""
    out: list[str] = []
    try:
        with zipfile.ZipFile(src) as z:
            for name in z.namelist():
                if not re.search(part_pattern, name):
                    continue
                xml = z.read(name).decode("utf-8", "replace")
                for para in re.findall(f"<{para_tag}[ >].*?</{para_tag}>", xml, re.S):
                    # **必須解 XML 實體**:原始檔裡的 `&` 是 `&amp;`,輸出
                    # 是 `&`,不解的話每個含 & < > 的段落都會誤報成「沒轉出來」
                    joined = html.unescape("".join(
                        re.findall(f"<{text_tag}[^>]*>([^<]*)</{text_tag}>", para),
                    ))
                    if joined.strip():
                        out.append(joined)
    except Exception:  # pragma: no cover - 稽核失敗絕不能拖垮轉檔
        logger.debug("自我稽核讀不到 %s", src, exc_info=True)
    return out


def missing_paragraphs(
    src: Path, blocks: list[Block], *, pptx: bool = False, xlsx: bool = False,
) -> list[str]:
    """原始檔裡有、但沒出現在輸出裡的段落文字。

    公開給 `scripts/audit_office.py` 共用——那支工具與這裡的自我稽核**必須
    是同一份比對邏輯**,否則會出現「稽核說有問題、實際輸出沒問題」這種
    最浪費時間的落差(實際發生過:稽核工具自己那份把整個群組的文字串起來
    比對,而輸出是逐形狀拆開的,一次掃描噴出 400 多筆誤報)。"""
    if pptx:
        originals = _xml_paragraphs(src, r"ppt/slides/slide\d+\.xml$", "a:p", "a:t")
    elif xlsx:
        # 儲存格文字有**兩種存法**:共用字串表(Excel 自己存的一律走這條)
        # 與行內字串(`<is><t>`,openpyxl 之類的產生器會這樣寫)。只認一種
        # 的話,另一種來源的檔案會整份比對不到、稽核形同虛設
        originals = _xml_paragraphs(src, r"xl/sharedStrings\.xml$", "si", "t")
        originals += _xml_paragraphs(src, r"xl/worksheets/sheet\d+\.xml$", "is", "t")
        originals += _xml_paragraphs(src, r"xl/drawings/drawing\d+\.xml$", "a:p", "a:t")
    else:
        originals = _xml_paragraphs(src, r"word/document\.xml$", "w:p", "w:t")
        # 頁首頁尾自 2026-08-02 起也會抽文字,對照就得跟著涵蓋——否則那條
        # 走訪路徑壞掉時稽核完全看不到。**過濾要用同一個判準**:輸出端
        # 丟掉的頁碼家具,對照端也得丟,不然「第 1 頁,共 25 頁」這種夠長的
        # 頁碼會被報成落差(政策差異不是走訪漏洞,一份出處才不會分岔)
        originals += [
            p for p in _xml_paragraphs(
                src, r"word/(header|footer)\d*\.xml$", "w:p", "w:t")
            if not _is_page_furniture(p)
        ]
        # 註腳同理(2026-08-04 起會抽,見 _note_blocks):內容在另一個 part,
        # 對照側不跟著涵蓋的話,那條走訪路徑壞掉時稽核完全看不到——而它
        # 「原本就沒抽」的那段日子裡,稽核也確實一聲不吭
        originals += _xml_paragraphs(
            src, r"word/(footnotes|endnotes)\.xml$", "w:p", "w:t")
    return docmd.missing_from(originals, blocks) if originals else []


def _extraction_gaps(src: Path, blocks: list[Block], *, pptx: bool) -> list[Block]:
    r"""**自我稽核**:原始檔裡的文字,有沒有全部出現在輸出裡?

    這是對「某一種 XML 節點沒被走訪」這類 bug 的常設防線。那種漏是**無聲**
    的——已經踩過三次(只認 DrawingML 的 a:blip、只走 body 直屬段落、
    文字方塊與群組整組沒走),而且沒辦法靠「多找幾個檔來看」窮舉
    (使用者 2026-08-02:「我隨手找兩個檔就發現問題,擔心還有沒被發現的」)。

    做法是拿 zip 裡的段落文字跟輸出比對,**刻意不經 python-docx/pptx**
    ——它們看不到的節點正是要抓的東西。比對只看「整段都不見」,並且把
    空白全部壓掉:輸出會重排、合併、加標點,逐字比對只會製造雜訊。

    抓到就在 md 裡留標記(進 frontmatter 的 lossy_kinds)。這是**工具自己
    的 bug**,但寧可讓使用者看見、也不要安靜地少——那正是本功能宣稱最不能
    接受的事,而且看得見才修得掉。"""
    missing = missing_paragraphs(src, blocks, pptx=pptx)
    if missing:
        logger.warning("%s:有 %d 段文字沒有被轉出來", src.name, len(missing))
    return docmd.extraction_gap_note(missing)


def _body_children(element, depth: int = 0):
    """依文件順序走出所有 `w:p` / `w:tbl`,**穿透內容控制項**(`w:sdt`)。

    Word 的封面頁、文件屬性欄位、目錄、表單控制項都包在 `w:sdt` 裡,實際
    內容在它的 `w:sdtContent` 底下。只認 body 直屬的 `w:p`/`w:tbl` 的話,
    那些整組看不到——實際案例(使用者 2026-08-02 的自我稽核):一份提案書
    的**文件標題**「台恆幣(TWDT-ETH)發行計劃書」與封面上的另外 11 段
    文字全部消失,而那是整份文件最重要的一行。"""
    for child in element.iterchildren():
        tag = str(child.tag).split("}")[-1]
        if tag in ("p", "tbl"):
            yield child
        elif tag == "sdt" and depth < _MAX_SDT_DEPTH:
            for content in child.iterchildren():
                if str(content.tag).split("}")[-1] == "sdtContent":
                    yield from _body_children(content, depth + 1)


# 頁碼那一類「版面家具」:去掉數字、空白與這些符號之後就沒剩什麼的,
# 對 RAG 毫無價值。實測 150 份真實公司文件有 58 份有頁首頁尾,其中純頁碼
# 佔了一半上下(「1」「第2頁」「- 7 -」「1 ①－」)
_FURNITURE_RE = re.compile(r"[\d\s第頁共之/／\-–—.,()\[\]①-⑳－&]+")


def _is_page_furniture(text: str) -> bool:
    """這段頁首頁尾文字是不是只有頁碼?

    判準是「拿掉數字與頁碼符號後還剩幾個字」而不是白名單:表單編號
    (`ADM-F-13`)、文件名(`TWDT-ETH發行計劃書12`)都要留,而它們拿掉
    數字之後仍有實質內容。代價是footer 裡孤零零一個日期(`7/25/2010`)
    會被當成家具丟掉——那種也確實不是內容。"""
    return len(_FURNITURE_RE.sub("", text or "")) < 2


def _header_footer_blocks(document) -> list[Block]:
    """頁首頁尾的文字。**每頁重複,所以只在整份文件的最前面收一次**。

    公司範本很愛把**文件標題、密級標示、版本編號**放在頁首頁尾,那些對
    RAG 是高價值的 metadata(使用者 2026-08-02 指定要抽)。但它是每頁都
    重複的內容:照頁輸出的話一份 30 頁的文件會多出 30 份一模一樣的字,
    切塊之後每個 chunk 都被同一句話污染。所以整份去重、只留一份。

    一節可以有三種頁首(首頁/奇數頁/偶數頁),多節文件又各有一組,而它們
    的內容多半相同——去重的鍵是壓過空白的文字,不是來源 part。

    **只取文字不取圖**:頁首裡的圖幾乎一律是公司信紙 logo,每頁一張、
    對內容毫無貢獻。純數字(頁碼)也不要——`docmd.GAP_MIN_CHARS` 那條
    長度門檻順便擋掉了,不必另外判斷。

    加了前綴讓區塊能獨立理解(docmd.render 規則 3):切塊器把它切出去
    之後,「機密」兩個字本身看不出是誰的機密。"""
    from docx.table import Table as DocxTable

    out: list[Block] = []
    seen: set[str] = set()
    for section in getattr(document, "sections", []) or []:
        for attr, label in (
            ("first_page_header", "頁首"), ("even_page_header", "頁首"),
            ("header", "頁首"),
            ("first_page_footer", "頁尾"), ("even_page_footer", "頁尾"),
            ("footer", "頁尾"),
        ):
            part = getattr(section, attr, None)
            element = getattr(part, "_element", None)
            if element is None:
                continue
            # 頁首頁尾各自是獨立的 part,rels 也是各自一份(見 _external_links)
            links = _external_links(getattr(part, "part", None))
            pieces: list[str] = []
            for child in _body_children(element):
                if str(child.tag).split("}")[-1] == "tbl":
                    rows, _ = _docx_table_rows(DocxTable(child, part))
                    # **列內去重**:頁首幾乎都用一個跨欄的版面表格排版,而
                    # `row.cells` 對合併儲存格是每個網格欄各給一份(那在真正的
                    # 表格裡是刻意的,見 KIND_MERGED_CELLS)。攤成一行文字之後
                    # 就變成「國立政治大學 商學院 國立政治大學 商學院」
                    pieces.extend(" ".join(dict.fromkeys(r)) for r in rows)
                else:
                    pieces.append(_para_text(child, links))
                    pieces.extend(_textbox_texts(child))
            text = _clean(" ".join(p for p in pieces if p.strip()))
            key = docmd.squash(text)
            if not key or key in seen or _is_page_furniture(text):
                continue
            seen.add(key)
            out.append(Para(f"{label}:{text}"))
    return out


# 註腳與章節附註:(rels 的型別尾巴, 節點名, 內文記號, 章節標題)
_NOTE_PARTS = (
    ("footnotes", "footnote", docmd.FOOTNOTE_MARK, "註腳"),
    ("endnotes", "endnote", docmd.ENDNOTE_MARK, "章節附註"),
)


def _related_part(part, tail: str):
    """依 rels 的型別找出關聯的 part(內部關聯,不是外部網址)。"""
    for rel in getattr(getattr(part, "rels", None), "values", lambda: [])():
        if getattr(rel, "is_external", False):
            continue
        if str(getattr(rel, "reltype", "")).endswith(f"/{tail}"):
            return getattr(rel, "target_part", None)
    return None


def _note_blocks(document) -> list[Block]:
    """註腳與章節附註,收在正文之後。

    **python-docx 完全不暴露它們**:內容在另一個 part(`word/footnotes.xml`),
    `document.paragraphs` 只走 body,所以整組安靜地不見。更糟的是**自我
    稽核也看不到**——`missing_paragraphs` 原本只讀 `word/document.xml`,
    對照側跟輸出側一起瞎,那正是這個專案最不能接受的那種漏(2026-08-04
    參考 markitdown 的 docx 前處理時發現;它連 footnotes/endnotes 都會
    一起前處理)。實測 424 份真實 docx 有 8 份(1.9%)有註腳、0 份有章節
    附註,量不大,但一份合約的註腳往往正是那份合約真正的條件。

    **分隔線那兩筆要跳掉**:每份 footnotes.xml 都有 `w:type="separator"`
    與 `continuationSeparator` 兩個內建項目(就是頁尾那條橫線),它們沒有
    文字,留著只會變成兩個空區塊。

    編號用檔案裡的 `w:id` 而不是自己從 1 數:正文的記號(`_para_text` 產出
    的 `[註3]`)也是用同一個 id,兩邊天然對得起來,不必再維護一份對照。"""
    from docx.oxml import parse_xml

    out: list[Block] = []
    for tail, tag, mark, label in _NOTE_PARTS:
        part = _related_part(getattr(document, "part", None), tail)
        blob = getattr(part, "blob", None)
        if not blob:
            continue
        try:
            root = parse_xml(blob)
        except Exception:  # noqa: BLE001 - 註腳讀不動不該讓整份文件陣亡
            logger.debug("讀不到 %s,略過", tail, exc_info=True)
            continue
        links = _external_links(part)
        items: list[Block] = []
        for node in root.iter(f"{_NS_W}{tag}"):
            if node.get(f"{_NS_W}type"):  # separator / continuationSeparator
                continue
            wid = node.get(f"{_NS_W}id")
            text = _clean("\n".join(
                _para_text(p, links) for p in node.iter(f"{_NS_W}p")
            ))
            if not text or not wid:
                continue
            # 前綴讓區塊能獨立理解(docmd.render 規則 3):切塊之後
            # 「見前揭註」四個字本身看不出是誰的註腳
            items.append(Para(f"{docmd.note_marker(mark, wid)} {text}"))
        if items:
            out.append(Heading(1, label))
            out.extend(items)
    return out


def _numbered_paragraphs(document, body) -> int:
    """帶 Word 自動編號的段落數。**直接掛的與樣式帶來的都要算**。

    按編號鈕會在段落上直接掛 `w:numPr`,但套用「清單段落」這類樣式時
    numPr 在 styles.xml 的樣式定義裡、段落本身乾乾淨淨(實測 python-docx
    預設範本的 `List Number` 就是這樣)——只認前者的話,整份用樣式排清單的
    文件會被判成「沒有編號」,而這個標記存在的理由正是不要有那種假乾淨。"""
    numbered_styles = set()
    for style in document.styles:
        element = getattr(style, "element", None)
        if element is not None and element.find(f".//{_NS_W}numPr") is not None:
            numbered_styles.add(style.style_id)
    n = 0
    for para in body.iter(f"{_NS_W}p"):
        ppr = para.find(f"{_NS_W}pPr")
        if ppr is None:
            continue
        if ppr.find(f"{_NS_W}numPr") is not None:
            n += 1
            continue
        pstyle = ppr.find(f"{_NS_W}pStyle")
        if pstyle is not None and pstyle.get(f"{_NS_W}val") in numbered_styles:
            n += 1
    return n


def convert_docx(src: Path, assets: AssetsDir, ocr_enabled: bool = True) -> list[Block]:
    """docx → Blocks(保住段落與表格的**原始先後順序**)。"""
    mod = _ensure_docx()
    try:
        document = mod.Document(str(src))
    except Exception as e:
        raise UserFacingError(
            f"無法開啟 Word 檔「{src.name}」:檔案可能已損壞,或其實是舊版 .doc 格式"
            "(請用 Word 另存為 .docx)"
        ) from e

    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph

    blocks: list[Block] = []
    body = document.element.body

    # 追蹤修訂:安靜地採用「最終版」是嚴重的無聲失真——審閱中的文件裡,
    # 刪除線的舊文字與新文字意思可能完全相反
    if body.find(f".//{_NS_W}ins") is not None or body.find(f".//{_NS_W}del") is not None:
        blocks.append(Note(
            "本文件含追蹤修訂,以下內容是「接受所有修訂」後的最終版",
            docmd.KIND_TRACKED_CHANGES,
        ))

    # 自動編號:編號本身不打算撈(見 docmd.KIND_NUMBERING),但**不能不說**
    # ——一份 76% 段落帶編號的文件輸出 lossy: 0,那是主動宣稱「什麼都沒丟」
    n_numbered = _numbered_paragraphs(document, body)
    if n_numbered:
        blocks.append(Note(
            f"本文件有 {n_numbered} 段使用 Word 的自動編號(數字清單或項目符號),"
            "編號與清單階層沒有呈現在文字裡",
            docmd.KIND_NUMBERING,
        ))

    # 頁首頁尾擺在最前面:它是整份文件的 metadata(標題/密級/版本),
    # 不屬於任何一頁的內文
    blocks.extend(_header_footer_blocks(document))

    links = _external_links(document.part)
    nested_seen = False
    for child in _body_children(body):
        tag = str(child.tag).split("}")[-1]
        if tag == "p":
            para = Paragraph(child, document)
            text = _clean(_para_text(child, links))
            level = _heading_level(para)
            if text:
                blocks.append(Heading(level, text) if level else Para(text))
            # 文字方塊的內容 para.text 取不到(見 _textbox_texts);同一段
            # 已經有的就不重複(有些方塊只是把本文再顯示一次)
            for boxed in _textbox_texts(child):
                if boxed not in text:
                    blocks.append(Para(boxed))
            # SmartArt 的內容在另一個 part,一般走訪看不到(見 _diagram_texts)
            for dgm in _diagram_texts(child, document.part):
                blocks.append(Para(f"SmartArt 圖形的文字:{dgm}"))
            blocks.extend(_para_images(para, document, assets, ocr_enabled))
        elif tag == "tbl":
            rows, nested = _docx_table_rows(DocxTable(child, document))
            nested_seen = nested_seen or nested
            if rows:
                blocks.append(Table(rows, has_header=True))
            # 表格儲存格裡的圖也要抽:版面用表格排版的文件(公司文件很常見)
            # 會把插圖整個放進儲存格,只走直屬段落的話一張都看不到
            blocks.extend(_images_from(child, document, assets, ocr_enabled))
            flat = "\n".join(" ".join(r) for r in rows)
            for boxed in _textbox_texts(child):
                if boxed not in flat:
                    blocks.append(Para(boxed))
            for dgm in _diagram_texts(child, document.part):
                blocks.append(Para(f"SmartArt 圖形的文字:{dgm}"))

    if nested_seen:
        blocks.append(Note(
            "本文件有表格中的表格,內層表格已攤平成同一格內的文字",
            docmd.KIND_NESTED_TABLE,
        ))
    # 註腳收在正文之後、巨集大綱之前:它是內文的一部分,不是檔案層級的附註
    blocks.extend(_note_blocks(document))
    blocks.extend(_macro_blocks(src))
    blocks.extend(_extraction_gaps(src, blocks, pptx=False))
    return blocks


# ---- pptx ----

def _shape_sort_key(shape):
    """形狀的閱讀順序:先上後下、再左至右。

    pptx 的 shapes 順序是「疊放順序」(z-order)不是閱讀順序——標題常常
    排在內文之後。top/left 對某些形狀是 None(例如置於版面配置佔位符),
    用大數推到最後而不是讓排序炸掉。"""
    top = getattr(shape, "top", None)
    left = getattr(shape, "left", None)
    return (top if top is not None else 1 << 30, left if left is not None else 1 << 30)


# 群組可以巢狀。上限只是防呆(壞檔造出來的自我參照),正常簡報一兩層
_MAX_GROUP_DEPTH = 8


def _flatten_shapes(shapes, depth: int = 0) -> list:
    """把群組攤平成一串形狀(遞迴)。

    `slide.shapes` 只給**頂層**形狀:群組(`p:grpSp`)裡的圖片、文字方塊、
    表格都在 `group.shapes` 底下,不遞迴的話整組看不到——實際案例
    (使用者 2026-08-02 的稽核):一份簡報 18 張圖有 **7 張在群組裡**,
    一張都沒抽到。攤平而不是「整組當一個形狀」是因為每個子形狀都有自己的
    位置,攤平後 `_shape_sort_key` 照樣排得出閱讀順序。"""
    out = []
    for shape in shapes:
        sub = getattr(shape, "shapes", None)
        if sub is not None and depth < _MAX_GROUP_DEPTH:
            out.extend(_flatten_shapes(sub, depth + 1))
        else:
            out.append(shape)
    return out


def _pptx_table_rows(table) -> list[list[str]]:
    return [[_clean(c.text) for c in row.cells] for row in table.rows]


def _defines_bullets(element) -> bool:
    """這個元素底下有沒有定義項目符號或自動編號?"""
    return element is not None and (
        element.find(f".//{_NS_A}buChar") is not None
        or element.find(f".//{_NS_A}buAutoNum") is not None
    )


def _master_defines_bullets(prs) -> bool:
    """母片的內文樣式有沒有定義項目符號?

    PowerPoint 的內文佔位符**預設就有項目符號**,而那個定義寫在母片的
    `p:txStyles/p:bodyStyle` 裡——投影片自己的 XML 什麼都不寫。"""
    return any(
        _defines_bullets(master.element.find(f"{_NS_P}txStyles"))
        for master in getattr(prs, "slide_masters", []) or []
    )


def _is_body_placeholder(shape) -> bool:
    """是不是「內文」佔位符?母片的預設項目符號只套在這種形狀上。"""
    return bool(getattr(shape, "is_placeholder", False)) and "TITLE" not in str(
        getattr(getattr(shape, "placeholder_format", None), "type", ""),
    )


def _frame_inherits_bullets(frame, placeholder: bool, master_bullets: bool) -> bool:
    """這個文字框的段落有沒有「沒寫但實際上有」的項目符號?

    **三種來源,只認段落自己寫的會留下 8.5% 的假乾淨**——實測 200 份真實
    簡報,有 17 份畫面上明明是清單(「App engineer / Systems Analysis
    Engineer / …」)卻一個 `buChar` 都沒寫。那正是 docx 那次「樣式帶來的
    編號」同一個陷阱。這裡處理其中兩種:形狀自己的 `a:lstStyle`,以及
    版面配置/母片(PowerPoint 的內文佔位符預設就有項目符號)。

    **要求「同一個文字框裡有兩段以上」**:少了這條,每個單行的標題方塊
    都會被當成一項清單。"""
    try:
        filled = sum(1 for p in frame.paragraphs if (p.text or "").strip())
    except Exception:  # noqa: BLE001 - 怪形狀不該拖垮整份簡報
        return False
    return filled >= 2 and (
        _defines_bullets(getattr(frame, "_txBody", None))
        or (master_bullets and placeholder)
    )


def _bullet_kind(para, inherits: bool) -> str | None:
    """這一段是不是條列?回 `"num"`(自動編號)、`"bullet"`,或 None。"""
    ppr = para._p.find(f"{_NS_A}pPr")
    if ppr is None:
        return "bullet" if inherits else None
    if ppr.find(f"{_NS_A}buNone") is not None:
        return None  # 明講不要項目符號
    if ppr.find(f"{_NS_A}buAutoNum") is not None:
        return "num"
    # lvl>=1(第二層以上)本身就代表清單階層,即使沒寫任何 buChar
    if ppr.find(f"{_NS_A}buChar") is not None or (ppr.get("lvl") or "0") != "0":
        return "bullet"
    return "bullet" if inherits else None


# PowerPoint 的大綱層級是 0..8
_MAX_LIST_LEVEL = 8


def _pptx_para_text(para, links: dict[str, str] | None = None) -> str:
    """pptx 段落文字,**含超連結網址**。

    **不能只走 `para.runs`**:軟換行 `<a:br/>` 與欄位 `<a:fld>`(投影片
    編號、日期)不是 run,只收 run 會把兩行黏成一行、也會漏掉欄位文字。
    `content_children` 正是 python-pptx 自己算 `para.text` 時走的那份順序,
    照它走就與原本的行為逐字相同,只多了連結。

    pptx 的連結掛在 run 的 `a:rPr/a:hlinkClick` 上(元素名與 docx 的
    `w:hyperlink` 不同,所以兩邊各有一份實作)。**只做 run 層級**:形狀
    層級的點擊動作(`p:cNvPr/a:hlinkClick`)實測兩份帶連結的簡報都是 0 個,
    而它沒有對應的文字可以掛。"""
    children = getattr(getattr(para, "_p", None), "content_children", None)
    if not children:
        return para.text
    out: list[str] = []
    for elm in children:
        text = elm.text or ""
        rid = ""
        if elm.tag == f"{_NS_A}r":
            for hl in elm.iter(f"{_NS_A}hlinkClick"):
                rid = hl.get(f"{_NS_R}id") or ""
                break
        out.append(_as_link(text, (links or {}).get(rid, "")) if rid else text)
    return "".join(out)


def _frame_markdown(frame, inherits: bool) -> str:
    """文字框 → markdown。**條列輸出成巢狀清單**,其餘段落原樣。

    階層來自 `paragraph.level`,縮排每層兩格。pptx 與 docx 在這裡分道揚鑣:
    Word 的編號在 `numbering.xml` 由 Word 自己算、真的拿不到,而 pptx 的
    層級就明寫在 `lvl` 屬性上——修得起來就不該只標記(使用者 2026-08-02)。

    **自動編號輸出實際序號而不是一律 `1.`**:這份 md 的讀者是 AI,一整排
    「1.」比 1. 2. 3. 難理解,而 markdown 渲染器本來就會自己重編號,寫實際
    序號沒有副作用。切回淺層時要把更深層的計數歸零,否則第二組子清單會
    從上一組的尾數接下去。"""
    lines: list[str] = []
    counters: dict[int, int] = {}
    prev_is_item = False
    # 連結表從 frame 自己的 part 取:投影片與備忘稿是不同的 part,各有
    # 一份 rels(見 _external_links),呼叫端不必為此多傳一個參數
    links = _external_links(getattr(frame, "part", None))
    for para in frame.paragraphs:
        text = _clean(_pptx_para_text(para, links))
        if not text:
            continue
        kind = _bullet_kind(para, inherits)
        if kind is None:
            counters.clear()
            # 清單與一般段落之間要空一行,否則 markdown 會把下一行
            # 當成清單項目的延續
            if prev_is_item:
                lines.append("")
            lines.append(text)
            prev_is_item = False
            continue
        level = min(max(getattr(para, "level", 0) or 0, 0), _MAX_LIST_LEVEL)
        for deeper in [lv for lv in counters if lv > level]:
            del counters[deeper]
        if kind == "num":
            counters[level] = counters.get(level, 0) + 1
            marker = f"{counters[level]}."
        else:
            marker = "-"
        if not prev_is_item and lines:
            lines.append("")
        prefix = "  " * level + marker + " "
        # 段落內的軟換行(`<a:br/>`)要縮到與內文對齊,否則第二行頂在
        # 第 0 欄——markdown 的 lazy continuation 雖然仍算同一項,但讀起來
        # 像是清單結束了
        lines.append(prefix + text.replace("\n", "\n" + " " * len(prefix)))
        prev_is_item = True
    return "\n".join(lines).strip()


def _lost_bullets(prs) -> int:
    """條列**還沒被保住**的段落數——只剩表格儲存格裡的。

    文字框已經輸出成巢狀清單(見 `_frame_markdown`),那些不算失真了;
    但 markdown 表格的一格塞不下多行清單,格子裡的條列仍然只能攤平。
    標記要跟著實際情況縮小,不然 lossy 會變成「反正都有標」的雜訊。"""
    n = 0
    for slide in prs.slides:
        for shape in _flatten_shapes(slide.shapes):
            if not getattr(shape, "has_table", False):
                continue
            try:
                cells = [c for row in shape.table.rows for c in row.cells]
            except Exception:  # noqa: BLE001 - 怪表格不該拖垮整份簡報
                logger.debug("表格儲存格讀取失敗", exc_info=True)
                continue
            for cell in cells:
                for para in cell.text_frame.paragraphs:
                    if (para.text or "").strip() and _bullet_kind(para, False):
                        n += 1
    return n


def _fill_blip_rids(slide) -> list[str]:
    r"""投影片上**不是 `p:pic` 形狀**的圖片(背景圖、形狀填滿)的 rId。

    `slide.shapes` 完全看不到它們:背景圖掛在 `p:cSld/p:bg/p:bgPr/a:blipFill`,
    形狀填滿掛在 `p:spPr/a:blipFill`——兩者都不是 `p:pic`,所以 `shape.image`
    一律是 None。實際案例(2026-08-03 全碟稽核):一份「部門業務簡介」的
    5 張投影片**整頁就是一張背景圖、`p:spTree` 是空的**,抽出來 0 張圖,
    等於整份簡報什麼都沒有。

    `p:pic` 底下的跳過:那些 `shape.image` 已經處理過,重複抽會讓同一張圖
    在 md 裡出現兩次。"""
    rids: list[str] = []
    for blip in slide.element.iter(f"{_NS_A}blip"):
        rid = blip.get(f"{_NS_R}embed")
        if not rid or rid in rids:
            continue
        rids.append(rid)
    return rids


def _slide_fill_images(
    slide, slide_no: int, assets: AssetsDir, seen: set[str], ocr_enabled: bool,
) -> list[Block]:
    """背景圖與形狀填滿 → Blocks(見 _fill_blip_rids)。

    **跨投影片去重**:整份簡報套同一張範本背景時,40 張投影片會各抽一次
    同一張圖。以 rId 對應到的實體 part 去重,只留第一次——那張圖仍然在
    md 裡看得到,而不會變成 40 份一樣的檔案。"""
    out: list[Block] = []
    for rid in _fill_blip_rids(slide):
        try:
            part = slide.part.related_part(rid)
            data = part.blob
            suffix = Path(str(getattr(part, "partname", "img.png"))).suffix or ".png"
        except (KeyError, AttributeError):
            logger.debug("投影片 %d 的填滿圖取不到 rId=%s", slide_no, rid)
            continue
        digest = hashlib.sha1(data).hexdigest()  # noqa: S324 - 只當識別碼
        if digest in seen:
            continue
        seen.add(digest)
        data, suffix = docimage.to_displayable(data, suffix)
        link = assets.add_bytes(data, suffix)
        if link:
            out.append(Image(link, f"投影片 {slide_no} 的背景/填滿圖片"))
        out.extend(docimage.ocr_image_bytes(
            data, f"投影片 {slide_no} 的背景圖片", ocr_enabled,
        ))
    return out


def convert_pptx(src: Path, assets: AssetsDir, ocr_enabled: bool = True) -> list[Block]:
    """pptx → Blocks(一張投影片一個章節,含表格、內嵌圖與備忘稿)。"""
    mod = _ensure_pptx()
    try:
        prs = mod.Presentation(str(src))
    except Exception as e:
        raise UserFacingError(
            f"無法開啟 PowerPoint 檔「{src.name}」:檔案可能已損壞,或其實是舊版 .ppt"
            "格式(請用 PowerPoint 另存為 .pptx)"
        ) from e

    blocks: list[Block] = []
    counts = {"chart": 0, "opaque": 0}
    fill_seen: set[str] = set()  # 背景圖跨投影片去重(見 _slide_fill_images)
    # 母片的內文樣式有沒有定義項目符號:每份簡報算一次,傳給每個形狀
    master_bullets = _master_defines_bullets(prs)
    for idx, slide in enumerate(prs.slides, start=1):
        # 投影片標題併進章節標題,而不是當成一般段落:RAG 切塊後的 chunk
        # 看到「投影片 7:第三季營收檢討」就知道自己在講什麼,只看
        # 「投影片 7」則毫無語意(docmd.render 規則 2)
        title_id = None
        title_text = ""
        try:
            title_shape = slide.shapes.title
            if title_shape is not None:
                title_text = _clean(getattr(title_shape, "text", ""))
                # 用 shape_id 而不是 `is` 比對:python-pptx 的 shapes 是即時
                # 產生的 proxy,每次存取都是**新的 Python 物件**,`is` 永遠
                # 不成立(踩過:標題會同時出現在章節標題與內文段落)
                title_id = getattr(title_shape, "shape_id", None)
        except Exception:
            logger.debug("投影片 %d 取不到標題", idx, exc_info=True)
        blocks.append(Heading(1, f"投影片 {idx}" + (f":{title_text}" if title_text else "")))

        for shape in sorted(_flatten_shapes(slide.shapes), key=_shape_sort_key):
            if title_id is not None and getattr(shape, "shape_id", None) == title_id:
                continue  # 已併進章節標題
            try:
                blocks.extend(_pptx_shape(
                    shape, idx, assets, counts, ocr_enabled, master_bullets,
                    fill_seen,
                ))
            except UserFacingError:
                raise
            except Exception:
                # 單一形狀讀不動不該讓整份簡報陣亡:python-pptx 對少見的
                # 形狀型別會拋 NotImplementedError,而一份 80 頁的簡報裡
                # 有一個怪形狀是常態
                logger.debug("投影片 %d 有形狀無法解析,略過", idx, exc_info=True)
                counts["opaque"] += 1

        # **形狀跑完才補撈**:背景圖、形狀填滿,以及 `mc:AlternateContent`
        # 裡 python-pptx 列不到的 p:pic。判準是「這份圖的位元組抽過沒有」
        # 而不是「它在不在 p:pic 底下」——後者會把包在 AlternateContent 裡
        # 的圖一起排除掉(實測一份簡報 64 個 blip 只抽到 24 個)
        blocks.extend(_slide_fill_images(
            slide, idx, assets, fill_seen, ocr_enabled,
        ))

        # SmartArt **掃整張投影片的 XML**,不跟著 shapes 走:它常被包在
        # `mc:AlternateContent` 裡,而 python-pptx 的 slide.shapes 根本不會
        # 列出那種元素——跟著 shape 找的話一個都抓不到(實測 0/5)
        for dgm in _diagram_texts(slide.element, slide.part):
            blocks.append(Para(f"SmartArt 圖形的文字:{dgm}"))

        notes = getattr(slide, "notes_slide", None)
        # 備忘稿也走同一條:講稿常常就是一份條列
        note_frame = getattr(notes, "notes_text_frame", None)
        note_text = _frame_markdown(note_frame, False) if note_frame is not None else ""
        if note_text:
            blocks.append(Para(f"**備忘稿**:{note_text}"))

    if counts["opaque"]:
        blocks.append(Note(
            f"有 {counts['opaque']} 個圖形(SmartArt、群組或特殊物件)的內容無法取出",
            docmd.KIND_SMARTART,
        ))
    if counts["chart"]:
        blocks.append(Note(
            f"本簡報有 {counts['chart']} 個圖表讀不到底層資料,只留下標註",
            docmd.KIND_CHART,
        ))
    # 文字框的條列已經輸出成巢狀清單,只剩表格格子裡的攤平了
    n_bullets = _lost_bullets(prs)
    if n_bullets:
        blocks.append(Note(
            f"表格儲存格裡有 {n_bullets} 段是條列,markdown 的表格一格塞不下"
            "多行清單,已攤平成同一格的文字",
            docmd.KIND_NUMBERING,
        ))
    blocks.extend(_macro_blocks(src))
    blocks.extend(_extraction_gaps(src, blocks, pptx=True))
    return blocks


# 佔比只對「整體的組成」有意義。長條/折線/散佈圖加一欄佔比是憑空捏造的
# 語意(那些圖的數列彼此獨立,加總沒有意義)
_PART_OF_WHOLE = ("PIE", "DOUGHNUT")


def _chart_title(chart) -> str:
    """圖表標題。取不到明文時退回**單一數列的名稱**。

    PowerPoint 對只有一個數列的圖,預設拿那個數列名當標題顯示,但 XML 裡的
    `<c:title>` 沒有明文 runs——`chart_title.text_frame.text` 因此回空字串,
    而畫面上明明寫著「Total」(使用者 2026-08-02 回報)。"""
    try:
        if chart.has_title:
            text = _clean(chart.chart_title.text_frame.text)
            if text:
                return text
    except Exception:
        logger.debug("圖表標題讀取失敗", exc_info=True)
    try:
        series = list(chart.plots[0].series)
        if len(series) == 1:
            return _clean(series[0].name or "")
    except Exception:
        logger.debug("圖表數列名稱讀取失敗", exc_info=True)
    return ""


def _num(value) -> str:
    """圖表數值 → 字串。整數不要拖著 `.0`(433.0 讀起來像有小數精度)。"""
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value)


def _pptx_chart(chart, slide_no: int) -> list[Block]:
    """圖表 → **把底層資料轉成表格**,而不是只留一句「有圖表」。

    資料就內嵌在 `ppt/charts/chartN.xml` 裡,python-pptx 的
    `plot.categories` 與 `series.values` 直接讀得到——原本只做計數、抓標題
    就回一個 Note,等於把整張圖的內容丟掉(使用者 2026-08-02 回報:一張
    環圈圖的 Developing 433 / Document 637 全沒了,而那是簡報的重點數字)。
    這比 xlsx 那邊還容易:xlsx 只拿得到 `'明細'!$B$2:$B$99` 這種參照,
    pptx 的數字就在檔案裡。

    圓餅/環圈另外加一欄佔比——那種圖表達的就是「佔整體多少」,而讀者
    (與 AI)從 433 這個數字看不出 40.5%。其他型別不加:長條/折線/散佈圖
    的數列彼此獨立,加總沒有意義,擅自算佔比是憑空捏造語意。"""
    title = _chart_title(chart)
    label = f"投影片 {slide_no} 的圖表" + (f":{title}" if title else "")
    kind = str(getattr(chart, "chart_type", "") or "")
    try:
        plot = chart.plots[0]
        categories = [_clean(str(c)) for c in plot.categories]
        series = [(_clean(s.name or ""), list(s.values)) for s in plot.series]
    except Exception:
        logger.debug("投影片 %d 的圖表資料讀不到", slide_no, exc_info=True)
        series = []
    if not series or not any(vals for _, vals in series):
        return [Note(
            f"此處有圖表{(':' + title) if title else ''},"
            "但讀不到底層資料,Markdown 無法呈現",
            docmd.KIND_CHART,
        )]

    share = any(k in kind.upper() for k in _PART_OF_WHOLE) and len(series) == 1
    header = ["項目", *(name or "數值" for name, _ in series)]
    if share:
        header.append("佔比")
    total = sum(v for v in series[0][1] if isinstance(v, (int, float))) if share else 0
    rows = [header]
    for i, cat in enumerate(categories or [""] * len(series[0][1])):
        row = [cat]
        for _, vals in series:
            row.append(_num(vals[i]) if i < len(vals) else "")
        if share:
            v = series[0][1][i] if i < len(series[0][1]) else None
            row.append(
                f"{v / total * 100:.1f}%"
                if isinstance(v, (int, float)) and total else ""
            )
        rows.append(row)
    caption = label + (f"({_CHART_KIND_ZH.get(kind.split()[0], '')})" if kind else "")
    return [Table(rows, has_header=True, caption=caption.replace("()", ""))]


def _pptx_shape(
    shape, slide_no: int, assets: AssetsDir, counts: dict, ocr_enabled: bool = True,
    master_bullets: bool = False,
    seen: set[str] | None = None,
) -> list[Block]:
    """單一形狀 → Blocks。判斷順序即優先序:圖表 → 表格 → 圖片 → 文字。

    `master_bullets` = 這份簡報的母片有沒有替內文佔位符定義項目符號
    (從 prs 才看得到,所以由呼叫端算好傳進來)。"""
    if getattr(shape, "has_chart", False):
        out = _pptx_chart(shape.chart, slide_no)
        # 只數「真的沒轉成文字」的:資料抽得出來時它已經是一張表格,
        # 再說一句「未能轉成文字」是假的失真標記,會讓 frontmatter 的
        # lossy 數字失去意義(RAG 建庫端拿它排序/設門檻)
        if any(isinstance(b, Note) for b in out):
            counts["chart"] += 1
        return out
    if getattr(shape, "has_table", False):
        rows = _pptx_table_rows(shape.table)
        # caption 讓表格被切塊後仍知道出處(docmd.render 規則 2)
        return [Table(rows, has_header=True, caption=f"投影片 {slide_no} 的表格")] if rows else []
    image = getattr(shape, "image", None)
    if image is not None:
        data = image.blob
        if seen is not None:
            seen.add(hashlib.sha1(data).hexdigest())  # noqa: S324 - 只當識別碼
        link = assets.add_bytes(data, f".{image.ext or 'png'}")
        out = [Image(link, f"投影片 {slide_no} 的圖片")] if link else []
        out.extend(docimage.ocr_image_bytes(
            data, f"投影片 {slide_no} 的圖片", ocr_enabled,
        ))
        return out
    if getattr(shape, "has_text_frame", False):
        frame = shape.text_frame
        inherits = _frame_inherits_bullets(
            frame, _is_body_placeholder(shape), master_bullets,
        )
        text = _frame_markdown(frame, inherits)
        return [Para(text)] if text else []
    return []


# ---- xlsx ----

def _merged_ranges(ws) -> list:
    """工作表的合併範圍。

    防呆集中在這裡:`ws` 可能是測試的假物件、或沒有 `merged_cells` 的
    read_only 變體。兩個呼叫端本來各寫了一種 getattr 拼法,防的是同一件事
    ——測試 monkeypatch 只會踩到其中一種,另一種等到真檔案才爆。"""
    return list(getattr(getattr(ws, "merged_cells", None), "ranges", None) or [])


def _merge_map(ws) -> tuple[dict[tuple[int, int], object], int]:
    """合併儲存格 → {(列, 欄): 左上角的值},以及合併範圍數。

    markdown 表格沒有 rowspan/colspan,只能把值展開重複到範圍內每一格
    ——這一定失真(讀的人分不出「本來就重複」與「本來是合併」),所以
    呼叫端一定要下 Note。"""
    filled: dict[tuple[int, int], object] = {}
    ranges = _merged_ranges(ws)
    for rng in ranges:
        span = (rng.max_row - rng.min_row + 1) * (rng.max_col - rng.min_col + 1)
        if span > _MAX_MERGE_CELLS:
            logger.warning("合併範圍過大(%d 格),不展開:%s", span, rng)
            continue
        value = ws.cell(rng.min_row, rng.min_col).value
        if value is None:
            # 排版用的空白合併很常見,展開了 get 回來還是 None:純白工。
            # 50 個各 10 萬格的空白合併照展開是 GB 級的 dict
            continue
        for r in range(rng.min_row, rng.max_row + 1):
            for c in range(rng.min_col, rng.max_col + 1):
                filled[(r, c)] = value
    return filled, len(ranges)


def _is_coloured(cell) -> bool:
    """儲存格有沒有帶「顏色語意」(底色或非黑字色)。

    公司的 Excel 很常用紅底/黃底表示狀態,轉成 md 之後那層意思整個消失
    ——AI 不知道紅色代表逾期,但至少該知道「這裡本來有顏色標註」。"""
    fill = getattr(cell, "fill", None)
    if fill is not None and getattr(fill, "patternType", None):
        rgb = getattr(getattr(fill, "fgColor", None), "rgb", None)
        if isinstance(rgb, str) and rgb not in ("00000000", "FFFFFFFF"):
            return True
    font = getattr(cell, "font", None)
    rgb = getattr(getattr(font, "color", None), "rgb", None)
    return isinstance(rgb, str) and rgb not in ("FF000000", "00000000")


def _fmt(value) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value)


def _chart_notes(ws) -> list[Block]:
    """圖表 → 先試著把底層資料的位置講出來,再退回單純標註。

    openpyxl 讀得到 chart 的 title 與 series 參照(像 `'明細'!$B$2:$B$99`)。
    對 AI 來說「這張圖畫的是明細表 B 欄」遠比「這裡有一張圖」有用,而且
    那些數字通常就在同一份檔案的其他工作表裡、已經轉成文字了。"""
    out: list[Block] = []
    for chart in getattr(ws, "_charts", None) or []:
        title = ""
        try:
            title = _clean(chart.title.tx.rich.p[0].r[0].t)
        except Exception:
            pass
        refs: list[str] = []
        for ser in getattr(chart, "series", None) or []:
            ref = getattr(getattr(getattr(ser, "val", None), "numRef", None), "f", None)
            if ref:
                refs.append(str(ref))
        detail = f":{title}" if title else ""
        source = f",資料來源 {'、'.join(refs[:5])}" if refs else ""
        out.append(Note(
            f"此處有圖表{detail}{source},Markdown 無法呈現圖形",
            docmd.KIND_CHART,
        ))
    return out


# ---- 並排版面(為列印排的寬表)----
#
# 為列印排版的表很常把數個獨立區塊**左右並排**:2026-08-19 的實例是一份
# 公司分機表,24 欄其實是 6 個部門各佔 3~6 欄,「這個人屬於哪一部門」只靠
# 版面位置表達、欄位名裡一個字都沒有。照列讀會把不同區塊的內容綁成同一
# 筆——那份轉出來每一「筆」都混了 6 個部門的人,而且錯得看不出來(使用者
# 只好自己用 xlrd 逐格取座標還原)。
#
# 線索是**區塊標題的水平合併**:標題橫跨整個區塊,它的左右邊界就是區塊的
# 邊界,而且同一組邊界會在不同列反覆出現(每個部門一次)。反過來,一般表格
# 的合併集中在最上面幾列(多層表頭),不會散在整張表。

_LAYOUT_MIN_TITLE_ROWS = 2   # 同一組欄邊界至少要重複出現在幾列
_LAYOUT_MIN_SIDE_ROWS = 2    # 至少幾列「同一列並排 >=2 個標題」——並排的鐵證
_LAYOUT_SPREAD = 0.2         # 區塊標題要橫跨整張表的比例(排除多層表頭)
_LAYOUT_MIN_SOLO = 0.3       # 至少幾成的標題列上「只有一個」區塊標題
_LAYOUT_MIN_FILLED = 0.5     # 至少幾成的區塊標題,底下真的有屬於它的資料列
_LAYOUT_MAX_ORPHAN = 0.2     # 最多幾成的格子可以不屬於任何區塊


def _title_spans(ws, first_row: int, last_row: int) -> dict[tuple[int, int], set[int]]:
    """區塊標題的水平合併 → {(起欄, 迄欄): 出現在哪些列}。

    只認「單列高、有值」的水平合併:區塊標題就長這樣。垂直合併是儲存格
    的縱向跨列(常見於「同一個部門的多列共用一格」),與版面分欄無關。"""
    spans: dict[tuple[int, int], set[int]] = {}
    ranges = _merged_ranges(ws)
    for rng in ranges:
        if rng.max_col <= rng.min_col or rng.max_row != rng.min_row:
            continue
        if not first_row <= rng.min_row <= last_row:
            continue
        value = ws.cell(rng.min_row, rng.min_col).value
        if value is None or not str(value).strip():
            continue
        spans.setdefault((rng.min_col, rng.max_col), set()).add(rng.min_row)
    return spans


def _is_side_by_side(spans: dict[tuple[int, int], set[int]], extent: int) -> bool:
    """這張表是不是並排版面?

    這裡是**前四關**(每一關都擋掉一整類的誤判,拿使用者的 205 個真實工作表
    校準):同一組欄邊界**反覆出現**、同一列上**並排**好幾個標題、標題
    **分散在整張表**、各區塊的標題**不是永遠成對出現在同一列**。

    另外兩關**要還原完才判得出來**,在 `_layout_blocks` 裡:區塊底下真的
    有資料列、以及沒有一大票格子無家可歸。"""
    cand = {k: v for k, v in spans.items() if len(v) >= _LAYOUT_MIN_TITLE_ROWS}
    chosen: list[tuple[int, int]] = []
    # 重疊時取「標題列數多」的:版面中途換配置時會有跨好幾群的大標題
    # (分機表第 45 列的 A:L),不能讓那種零星的標題吃掉主結構
    for span in sorted(cand, key=lambda k: (-len(cand[k]), k)):
        if any(not (span[1] < c[0] or span[0] > c[1]) for c in chosen):
            continue
        chosen.append(span)
    # 並排的鐵證:同一列上同時有 >=2 個區塊的標題。⚠️ 這條**順帶保證了
    # 「至少有兩組欄邊界」**(同一列要放得下兩個標題,就得有兩組)——原本
    # 另外寫的 `_LAYOUT_MIN_GROUPS` 是死的,2026-08-19 code review 指出它
    # 沒有任何測試守得住,查下去才發現根本推不動:拿掉它結果一模一樣。⚠️ 這條與下面兩關在
    # 使用者那 205 個真實工作表上**結果重疊**(拿掉它命中的還是同樣兩個),
    # 留著是因為代價不對稱:切錯的失敗方式是安靜的,而它擋的是「左右兩疊
    # 標題從不對齊」這種其實不確定該不該切的版面
    per_row: dict[int, int] = {}
    for span in chosen:
        for r in cand[span]:
            per_row[r] = per_row.get(r, 0) + 1
    if sum(1 for n in per_row.values() if n >= 2) < _LAYOUT_MIN_SIDE_ROWS:
        return False
    # 各區塊的標題若**永遠成對出現在同一列**,那是「每隔幾十列重印一次的
    # 多層表頭」不是並排區塊(2026-08-19 實測誤傷:一份權限表的「檔案存取
    # 權限」E:H 與「目錄存取權限」I:M 成對出現 11 次,前三關全被騙過,切出
    # 來的東西比不切還糟)。真的並排時,各區塊各自換標題、很少對齊
    if sum(1 for n in per_row.values() if n == 1) < len(per_row) * _LAYOUT_MIN_SOLO:
        return False
    # 標題要分散在整張表:多層表頭的合併全擠在最上面那幾列,那是表頭不是
    # 並排區塊。⚠️ extent 是「**實際有內容的列數**」不是列號範圍:分機表的
    # 標題橫跨第 1~52 列、內容只有 65 列,但第 280 列有一格孤兒儲存格,拿
    # 列號範圍當分母(280)會讓真正的並排表過不了這一關
    rows = {r for span in chosen for r in cand[span]}
    return (max(rows) - min(rows)) >= extent * _LAYOUT_SPREAD


def _detect_side_by_side(ws, rows: list[list[str]]) -> dict[tuple[int, int], set[int]] | None:
    """並排版面的偵測**單一入口**:是的話回區塊標題的合併範圍,不是回 None。

    ⚠️ 只能有這一個入口。huge 模式為了記憶體不做還原、但仍要示警,本來在
    那條路上自己抄了一份偵測,`extent` 卻傳成 `len(rows)`(含中段空列)而不是
    「實際有內容的列數」——兩邊分母不同,同一張表在兩條路徑上可能得到相反
    的判定,而空白列多得離譜正是這種檔案的特徵(那份分機表 279 列裡 215 列
    全空),於是警告在最需要它的檔案上最容易失靈。2026-08-20 /simplify 抓到。"""
    last_row = len(rows)
    max_col = max((len(r) for r in rows), default=0)
    if last_row < 5 or max_col < 4:
        return None
    spans = _title_spans(ws, 1, last_row)
    extent = sum(1 for r in rows if docmd.row_has_content(r))
    return spans if _is_side_by_side(spans, extent) else None


def _layout_blocks(ws, rows: list[list[str]]) -> list[Block]:
    """並排版面 → 一個區塊一小節;不是並排版面就回 []。

    歸屬規則只有一條:**一格屬於「涵蓋它、且離它最近的上方標題」**。

    ⚠️ 不可改用「先把工作表切成幾個垂直欄群」那種做法(2026-08-19 試過,
    28 格無聲消失):版面會中途換配置——分機表第 45 列的「數位增長部」佔
    A:L,一口氣跨掉左邊三個欄群,但**同一列 M 欄以右另有六個人**。以列為
    單位分段就會把那些人連同第 1 列右上角的製表日期一起丟掉,而欄群模型
    無論怎麼補都補不回這種「標題只管自己那幾欄」的語意。"""
    spans = _detect_side_by_side(ws, rows)
    if spans is None:
        return []
    max_col = max((len(r) for r in rows), default=0)

    titles = sorted(
        (r, lo, hi, docmd.one_line(rows[r - 1][lo - 1]))
        for (lo, hi), rs in spans.items()
        for r in rs
    )
    first_title_row = titles[0][0]
    # 每一欄各自的標題,用來 O(log n) 找「上方最近的標題」。titles 已經排過
    # 序、又是照順序附加的,所以每一串天生遞增,不必再 sort
    by_col: dict[int, list[tuple[int, int]]] = {}
    # 標題自己佔的格不是資料(合併展開後整段都是標題文字)
    is_title_cell: set[tuple[int, int]] = set()
    for i, (r, lo, hi, _text) in enumerate(titles):
        for c in range(lo, hi + 1):
            by_col.setdefault(c, []).append((r, i))
            is_title_cell.add((r, c))

    # 標題 → 列 → 該區塊那幾欄的值。⚠️ 內層是**預先配好的 list** 不是 dict:
    # 一格一個 dict 項目在接近降級門檻的表上要多吃 90MB(8GB 基準機),而且
    # 輸出時本來就要攤成 list——實測換成 list 快 28%、記憶體少 3.4 倍
    owned: dict[int, dict[int, list[str]]] = {}
    orphan: dict[int, dict[int, str]] = {}             # 沒有標題涵蓋的格
    kept = orphan_cells = 0
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row, start=1):
            if not value.strip() or (r, c) in is_title_cell:
                continue
            kept += 1
            seq = by_col.get(c) or []
            pos = bisect.bisect_left(seq, (r, -1)) - 1
            if pos < 0:
                orphan.setdefault(r, {})[c] = value
                orphan_cells += 1
                continue
            i = seq[pos][1]
            by_row = owned.setdefault(i, {})
            line = by_row.get(r)
            if line is None:
                lo, hi = titles[i][1], titles[i][2]
                line = by_row[r] = [""] * (hi - lo + 1)
            line[c - titles[i][1]] = value

    # 最後兩道關卡,而且是**還原完才判得出來**的:
    # ① 表單版面(「標籤 | 值」左右並排的申請表)合併特徵與並排區塊幾乎
    #    一樣,但它合併起來的是**值**——底下沒有屬於它的明細。硬切會把標籤
    #    與值拆到不同小節、把明細表砍掉幾欄,比不切糟得多(2026-08-19 拿
    #    使用者的暫借款申請表實測)。區塊標題底下大多真的有人才算還原成功
    # ② 而且**不能有一大票格子無家可歸**:表單的「標籤」欄從頭到尾不被任何
    #    合併涵蓋,一混進來就是整整一欄落單;區塊標題比資料窄的表(電話欄
    #    凸出去一格)也一樣。兩種都會把人的資料從他自己那一列扯出來
    # 不合格就退回原本的整表
    if len(owned) < len(titles) * _LAYOUT_MIN_FILLED:
        return []
    if orphan_cells > kept * _LAYOUT_MAX_ORPHAN:
        return []

    # 有資格當上層的只有「第一列」那幾個標題(見下)。條件是迴圈不變量,
    # 先挑出來——放在迴圈裡就是 O(標題數²),實測 8000 個標題要 2.8 秒,
    # 全花在一個 H2/H3 的排版決定上(2026-08-20 /simplify 實測)
    top_titles = [t for t in titles if t[0] == first_title_row]

    out: list[Block] = []
    if orphan:
        # 不屬於任何區塊的格子。**一定要輸出**:轉不出來可以標記,安靜消失
        # 不行。⚠️ caption 不可以說成「區塊標題之前的內容」——它們未必在
        # 前面(整欄沒有標題的話,來源可能是表格中段)
        body = [
            [cells.get(c, "") for c in range(1, max_col + 1)]
            for _r, cells in sorted(orphan.items())
        ]
        out.append(Table(body, has_header=False, caption=f"{ws.title}(不屬於任何區塊的內容)"))
    for i, (row_no, lo, hi, text) in enumerate(titles):
        # 只有「**第一列**那個更寬的標題」才算上層(分機表:表名底下才是各
        # 部門)。⚠️ 不可以認任何更寬的標題:版面中段跨好幾個區塊的標題是
        # **兄弟**不是父母,認了會把後面的區塊都掛到它底下,宣告一個原檔
        # 根本沒表達的從屬關係——而 render 的規則 1 說「錯的階層比沒有階層
        # 更糟」(2026-08-19 code review 拿本檔自己的 fixture 重現)
        wider = row_no > first_title_row and any(
            o[1] <= lo and o[2] >= hi for o in top_titles
        )
        name = text or f"第 {lo}–{hi} 欄"
        out.append(Heading(3 if wider else 2, name))
        cells = owned.get(i)
        if not cells:
            continue   # 有標題、底下沒人:標題本身也是內容,留著就好
        body = [line for _rr, line in sorted(cells.items())]
        # caption 帶區塊名是**必要的**,不是重複(2026-08-19 code review):
        # 超過 12 欄的區塊會走 `render_records`,而那條路徑用 caption 當每
        # 一筆的前綴——沒有它,「### 第 N 筆」與區塊標題同級,切塊器會把那
        # 些人切成一個不知道屬於哪一部門的 chunk
        out.append(Table(body, has_header=False, caption=f"{ws.title} · {name}"))
    return out


def _sheet_blocks(ws, ws_formula, huge: bool) -> tuple[list[Block], bool]:
    """單一工作表 → (Blocks, 有沒有拆成並排區塊)。

    huge=True 時只讀值、跳過所有 metadata 檢查。⚠️ 「有沒有拆」要**由這裡
    回報**,不可以讓呼叫端回頭去嗅輸出裡的 `KIND_*`:那些 token 是給 RAG 端
    篩 `lossy_kinds` 用的,拿它當內部控制訊號的話,以後任何人為了多說一句話
    而增減一個 Note 都會遠端改到版面(huge 模式同樣會放 KIND_SIDE_BY_SIDE
    的警告 Note,卻**沒有**拆——2026-08-20 /simplify 抓到的實例)。"""
    blocks: list[Block] = [Heading(1, str(ws.title))]
    if huge:
        blocks.append(Note(
            "本工作表資料量龐大,已略過合併儲存格、底色與圖表的檢查以節省記憶體"
            "(資料本身完整,一列都沒有省略)",
            docmd.KIND_HUGE_SHEET,
        ))
        rows = [[_fmt(v) for v in row] for row in ws.iter_rows(values_only=True)]
    else:
        merges, n_merged = _merge_map(ws)
        coloured = 0
        rows = []
        for row in ws.iter_rows():
            line = []
            for cell in row:
                value = merges.get((cell.row, cell.column), cell.value)
                # has_style 短路:實測 2000 格只有 3 格為真,省 4 倍
                if cell.has_style and _is_coloured(cell):
                    coloured += 1
                line.append(_fmt(value))
            rows.append(line)
        if n_merged:
            blocks.append(Note(
                f"本工作表有 {n_merged} 處合併儲存格,已展開成一般儲存格",
                docmd.KIND_MERGED_CELLS,
            ))
        if coloured:
            blocks.append(Note(
                f"本工作表有 {coloured} 個儲存格以底色或字色標註,Markdown 無法呈現顏色",
                docmd.KIND_CELL_COLOR,
            ))
        blocks.extend(_chart_notes(ws))
        if getattr(ws, "_pivots", None):
            blocks.append(Note("此處有樞紐分析表,Markdown 無法呈現", docmd.KIND_PIVOT))
        cf = getattr(ws, "conditional_formatting", None)
        n_cf = len(list(cf)) if cf is not None else 0
        if n_cf:
            blocks.append(Note(f"本工作表有 {n_cf} 組條件式格式規則未呈現", docmd.KIND_COND_FORMAT))
        blocks.extend(_formula_blocks(ws, ws_formula, rows))

    # 尾端的整列空白剪掉:Excel 常把 max_row 撐得比實際資料遠
    while rows and not any(c.strip() for c in rows[-1]):
        rows.pop()
    if not rows:
        blocks.append(Note("這個工作表是空的", docmd.KIND_BLANK_PAGE, lossy=False))
        return blocks, False

    # 並排版面(為列印排的寬表)要先依區塊還原,否則同一列會把不同區塊的
    # 內容綁成一筆。
    # ⚠️ huge 模式跳過的理由是**記憶體**,不是「讀不到合併範圍」(2026-08-19
    # code review 更正:活頁簿在 huge 模式下同樣不是 read_only,合併範圍讀
    # 得到)——還原要為每個非空格建一筆歸屬,那正是 huge 模式在省的東西。
    # 但不能就這樣安靜地轉錯:偵測本身只看合併範圍、很便宜,照做,發現是
    # 並排版面就明講「這次沒有還原」
    layout = _layout_blocks(ws, rows) if not huge else []
    if huge and _detect_side_by_side(ws, rows) is not None:
        blocks.append(Note(
            "本工作表看起來是為列印排的並排版面(數個區塊左右並排),但資料量"
            "太大、這次沒有依區塊還原——**同一列的內容可能分屬不同區塊**,"
            "不要當成同一筆資料",
            docmd.KIND_SIDE_BY_SIDE,
        ))
    if layout:
        # 右往左掃、掃到目前最寬就收手:整片掃過去是一趟不會短路的 O(格數),
        # 只為了 Note 裡的一個數字(實測 230k 格 10.1ms → 0.14ms)
        width = 0
        for r in rows:
            for i in range(len(r) - 1, width - 1, -1):
                if r[i].strip():
                    width = i + 1
                    break
        blocks.append(Note(
            f"本工作表是為列印排的並排版面({width} 欄其實是數個左右並排的"
            "區塊),已依區塊還原成各自獨立的小節——原檔同一列的內容分屬"
            "不同區塊,不是同一筆資料",
            docmd.KIND_SIDE_BY_SIDE,
            lossy=False,
        ))
        blocks.extend(layout)
        return blocks, True

    # 中段的整列空白也剪掉(尾端那個 while 只咬得到最後面)。Excel 的
    # max_row 常被遠處一格沒清掉的舊資料撐開,中間就多出整片空列——留著的話
    # 表格是幾百列 `|  |  |`、寬表則是幾百個只有標題沒有內容的「第 N 筆」。
    # ⚠️ 要放在 `_formula_blocks` **之後**:那支拿 rows 當「第 N 列第 M 欄」
    # 的索引查快取值,先抽掉列會讓它對到別格
    rows = docmd.drop_blank_rows(rows)
    caption = f"{ws.title}(共 {len(rows)} 列)"
    # 跨欄的標題列(例如 A1:D1 合併成「第三季各區營收」)直接當表頭的話,
    # **真正的欄名會掉成第一列資料**——AI 之後引用欄位就全錯了。把它提升
    # 成表格說明,讓下一列當表頭。
    # ⚠️ 判準是「非空的格全都被水平合併涵蓋」,不是「整列同一個值」:標題
    # 列常常不只一段(2026-08-19 的分機表,左邊 A:U 是表名、右邊 V:X 是製表
    # 日期),兩個值就過不了同值那一關,於是 24 個欄名全變成表名、等於零
    # 資訊。反向保護還在:單欄表的第一列不在任何合併裡,不會被吃掉
    if len(rows) > 1:
        titles = _title_spans(ws, 1, 1)
        covered = {c for lo, hi in titles for c in range(lo, hi + 1)}
        filled = {i + 1 for i, c in enumerate(rows[0]) if c.strip()}
        # ⚠️ 還要求「有一段寬到佔掉半列以上」(2026-08-19 code review 補):
        # 只看「非空的格全被合併涵蓋」會誤殺**每個欄名各自合併兩欄**的表頭
        # ——那種第一列會被整列刪掉、第一筆資料被拱上去當欄名,比原本的毛病
        # 嚴重得多。真正的大標題是「少少幾段、每段很寬」
        widest = max((hi - lo + 1 for lo, hi in titles), default=0)
        if len(filled) > 1 and filled <= covered and widest * 2 > len(filled):
            names = [docmd.one_line(rows[0][lo - 1]) for lo, _hi in sorted(titles)]
            caption = f"{ws.title} — {'、'.join(names)}(共 {len(rows) - 1} 列)"
            rows = rows[1:]

    # 寬表要不要改逐筆區塊由 docmd.render 決定(渲染政策只有一個出處)
    blocks.append(Table(rows, has_header=True, caption=caption))
    return blocks, False


def _formula_blocks(ws, ws_formula, rows: list[list[str]]) -> list[Block]:
    """公式沒有快取值時的補救。

    openpyxl 的 `data_only=True` 取的是「Excel 上次存檔時算好並存進檔案的
    快取值」——**檔案必須真的被 Excel 開啟並存過**才有。程式產生的 xlsx
    (各種系統匯出的報表都是)讀出來整片是 None,使用者會看到一份空表卻
    不知道為什麼。對策:同時用 data_only=False 讀公式本身,值空、公式非空
    就把公式列出來,並指路怎麼拿到真正的數字。"""
    if ws_formula is None:
        return []
    # 只留前 20 個字串、其餘用計數器:整片都是無快取公式的大表會蒐集出
    # 24 萬個字串(約 20MB),而顯示只用得到 20 個
    samples: list[str] = []
    n_missing = 0
    for row in ws_formula.iter_rows():
        for cell in row:
            value = cell.value
            if not (isinstance(value, str) and value.startswith("=")):
                continue
            r, c = cell.row - 1, cell.column - 1
            cached = rows[r][c] if r < len(rows) and c < len(rows[r]) else ""
            if not cached:
                n_missing += 1
                if len(samples) < 20:
                    samples.append(f"{cell.coordinate}: {value}")
    if not n_missing:
        return []
    shown = "、".join(samples)
    more = f"(僅列出前 20 個,共 {n_missing} 個)" if n_missing > 20 else ""
    return [Note(
        f"本工作表有 {n_missing} 個公式沒有計算結果,檔案裡只存了公式本身"
        f"——請用 Excel 開啟後另存新檔再轉一次,就會有數字。公式:{shown}{more}",
        docmd.KIND_FORMULA_NO_CACHE,
    )]


def _sheet_drawing_texts(src: Path) -> dict[str, list[str]]:
    r"""工作表上**圖形/文字方塊裡的文字**,依工作表分組。

    **openpyxl 讀不到這些**:它管的是儲存格,而圖形的文字住在
    `xl/drawings/drawingN.xml`。機房配置圖、流程圖、示意圖這類 Excel 很常
    把內容全放在圖形上——實際案例(使用者 2026-08-02 的稽核):一份機櫃
    配置圖的機櫃編號有兩個只存在圖形裡,儲存格完全沒有。

    對應關係要走三段 rels:workbook → sheet → drawing。任何一段接不上就
    把文字歸到「(未知工作表)」——**寧可位置不精確,也不要整段消失**。"""
    out: dict[str, list[str]] = {}
    try:
        with zipfile.ZipFile(src) as z:
            names = set(z.namelist())
            wb = z.read("xl/workbook.xml").decode("utf-8", "replace")
            wb_rels = z.read("xl/_rels/workbook.xml.rels").decode("utf-8", "replace")
            rid_to_part = dict(re.findall(
                r'Id="([^"]+)"[^>]*Target="([^"]+)"', wb_rels,
            ))
            for m in re.finditer(
                r'<sheet[^>]*name="([^"]*)"[^>]*r:id="([^"]+)"', wb,
            ):
                title, rid = html.unescape(m.group(1)), m.group(2)
                part = rid_to_part.get(rid, "").lstrip("/")
                sheet_part = part if part.startswith("xl/") else f"xl/{part}"
                rels = sheet_part.replace(
                    "worksheets/", "worksheets/_rels/") + ".rels"
                if rels not in names:
                    continue
                srels = z.read(rels).decode("utf-8", "replace")
                for target in re.findall(r'Target="([^"]*drawing\d+\.xml)"', srels):
                    dpart = "xl/" + target.replace("../", "")
                    if dpart not in names:
                        continue
                    xml = z.read(dpart).decode("utf-8", "replace")
                    texts = [
                        html.unescape(x) for x in re.findall(r"<a:t>([^<]*)</a:t>", xml)
                        if x.strip()
                    ]
                    if texts:
                        out.setdefault(title, []).extend(texts)
    except Exception:  # pragma: no cover - 讀不到就當沒有,絕不拖垮轉檔
        logger.debug("讀不到 %s 的圖形文字", src, exc_info=True)
    return out


def _has_macros(src: Path) -> bool:
    """活頁簿裡有沒有巨集(VBA)。

    `.xlsm` 在格式上就是 `.xlsx`,差別只是 zip 裡多一個 `xl/vbaProject.bin`
    ——所以兩者共用同一個 reader 是正確的、不是權宜之計。但那份 VBA 是
    **二進位 OLE 複合檔**,要抽出原始碼得另外接解析器;使用者 2026-08-02
    決定不抽、只留標記。多數 .xlsm 的價值在資料不在程式,而少數「工具型
    活頁簿」的重點就是那段巨集——那時至少要讓人知道「這裡少了東西」。"""
    try:
        with zipfile.ZipFile(src) as z:
            return any("vbaproject" in n.lower() for n in z.namelist())
    except Exception:  # pragma: no cover - 讀不到就當沒有
        logger.debug("讀不到 %s 的巨集資訊", src, exc_info=True)
        return False


# VBA 模組的樣板行:`Attribute VB_Name`、`VB_Base`、`VB_GlobalNameSpace` 等
# 是 VBA 自己的中繼資料,不是使用者寫的程式。全部只有這些 = 空模組
_VBA_ATTR_RE = re.compile(r"^\s*Attribute\s+VB_\w+\s*=", re.IGNORECASE)


def _clean_vba(code: str) -> str:
    """去掉模組層的 `Attribute VB_*` 樣板,留下真正的程式。

    那幾行是 VBA 匯出格式的產物,對讀的人(與 AI)沒有意義,而且每個模組
    都有——一份 26 個模組的活頁簿光樣板就佔掉可觀的篇幅。**Sub 裡面的
    `Attribute xxx.VB_Description` 留著**:那是巨集的說明文字,有內容。"""
    lines = (code or "").splitlines()
    body = []
    heading = True
    for line in lines:
        if heading and (not line.strip() or _VBA_ATTR_RE.match(line)):
            continue
        heading = False
        body.append(line)
    return "\n".join(body).strip()


# 程序宣告:Sub / Function / Property,含 Public/Private/Friend/Static 修飾
_VBA_PROC_RE = re.compile(
    r"^\s*(?:(Public|Private|Friend)\s+)?(?:Static\s+)?"
    r"(Sub|Function|Property\s+(?:Get|Let|Set))\s+([\w一-鿿]+)",
    re.IGNORECASE | re.MULTILINE,
)
# 巨集在 Excel「巨集」對話框裡顯示的說明,VBA 存成這個屬性
_VBA_DESC_RE = re.compile(
    r'^\s*Attribute\s+([\w一-鿿]+)\.VB_Description\s*=\s*"(.*)"',
    re.IGNORECASE | re.MULTILINE,
)
# 每個模組最多列幾行開頭註解:那通常是「這個模組在做什麼」
_VBA_HEAD_COMMENTS = 3


def _vba_outline(src: Path) -> list[Block]:
    r"""活頁簿裡的巨集 → **大綱**(模組、程序名、說明),不是原始碼。

    使用者 2026-08-02 的目的很明確:「除了知道有 VBA 之外,只是想大概知道
    VBA 用途」。原始碼放不進去——實測一份活頁簿就有 **19 萬字元**,那是
    整份 md 的好幾倍,而且對「它在做什麼」這個問題,程序名與說明的資訊
    密度高得多(大綱只有原始碼的 1~2%)。

    三個來源合起來就足以回答用途:
    - **程序名**:`Sub 取得資料()`、`Sub 更新報表()` ——中文命名的巨集尤其直白
    - **`VB_Description`**:Excel「巨集」對話框裡顯示的說明,VBA 存成屬性
    - **模組開頭的註解**:寫巨集的人自己留的說明

    要看完整程式碼就開原始檔——輸出裡會明講這件事,不是安靜地少。"""
    try:
        from oletools.olevba import VBA_Parser
    except ImportError:  # pragma: no cover - 相依缺失只影響巨集,不擋轉檔
        logger.debug("olevba 不可用,略過巨集大綱")
        return []
    parser = None
    try:
        parser = VBA_Parser(str(src))
        if not parser.detect_vba_macros():
            return []
        modules = [(name, code or "") for _, _, name, code in parser.extract_macros()]
    except Exception:
        logger.debug("讀取 %s 的巨集失敗", src.name, exc_info=True)
        return []
    finally:
        if parser is not None:
            try:
                parser.close()
            except Exception:  # pragma: no cover
                pass

    total_chars = 0
    sections: list[tuple[str, list[str], list[str]]] = []
    for name, code in modules:
        body = _clean_vba(code)
        total_chars += len(body)
        descs = dict(_VBA_DESC_RE.findall(code))
        procs = [
            f"`{kind.split()[0].title()} {proc}` "
            + (f"— {descs[proc]}" if descs.get(proc) else "")
            for _scope, kind, proc in _VBA_PROC_RE.findall(body)
        ]
        comments = []
        for line in body.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if not stripped.startswith("'"):
                break
            comments.append(stripped.lstrip("'").strip())
            if len(comments) >= _VBA_HEAD_COMMENTS:
                break
        if procs or comments:
            sections.append((name, [p.strip() for p in procs], comments))
    if not sections:
        return []

    proc_count = sum(len(p) for _, p, _ in sections)
    out: list[Block] = [
        Heading(1, "巨集(VBA)"),
        Para(
            f"本檔含巨集:{len(sections)} 個模組、{proc_count} 個程序,"
            f"程式碼約 {total_chars // 1000} 千字元。"
            "**以下只列出結構與說明,完整程式碼請開啟原始檔查看。**"
        ),
    ]
    for name, procs, comments in sections:
        out.append(Heading(2, name))
        if comments:
            out.append(Para("模組註解:" + " / ".join(comments)))
        if procs:
            out.append(Raw("\n".join(f"- {p}" for p in procs)))
    return out


def convert_xlsx(src: Path, assets: AssetsDir, ocr_enabled: bool = True) -> list[Block]:
    """xlsx/xlsm → Blocks(一個工作表一個章節)。"""
    mod = _ensure_openpyxl()
    try:
        wb = mod.load_workbook(str(src), data_only=True)
    except Exception as e:
        raise UserFacingError(
            f"無法開啟 Excel 檔「{src.name}」:檔案可能已損壞、有密碼保護,"
            "或其實是舊版 .xls 格式(請用 Excel 另存為 .xlsx)"
        ) from e

    drawing_texts = _sheet_drawing_texts(src)

    total_cells = sum(
        (ws.max_row or 0) * (ws.max_column or 0) for ws in wb.worksheets
    )
    huge = total_cells > _HUGE_CELLS
    wb_formula = None
    if not huge:
        try:
            # read_only:第二份只被 _formula_blocks 用來讀 value/coordinate,
            # ReadOnlyCell 全都有。實測 8000×30 的表:12.9s/97MB → 0.03s/1MB
            wb_formula = mod.load_workbook(
                str(src), data_only=False, read_only=True,
            )
        except Exception:
            logger.exception("讀取公式失敗,略過公式檢查:%s", src.name)

    blocks: list[Block] = []
    for ws in wb.worksheets:
        ws_formula = None
        if wb_formula is not None and ws.title in wb_formula.sheetnames:
            ws_formula = wb_formula[ws.title]
        sheet_blocks, split = _sheet_blocks(ws, ws_formula, huge)
        blocks.extend(sheet_blocks)
        extras: list[Block] = _sheet_images(ws, assets, ocr_enabled)
        # 圖形/文字方塊裡的字 openpyxl 讀不到(見 _sheet_drawing_texts)
        shape_text = _clean(" ".join(drawing_texts.get(ws.title, [])))
        if shape_text:
            extras.append(Para(f"工作表上圖形的文字:{shape_text}"))
        # 並排版面已經把工作表拆成一個個小節,這些**整張表**的附屬內容再
        # 直接接上去,就會變成「最後那一個區塊的東西」——機櫃圖正是重災區
        # (`_sheet_drawing_texts` 記著:有兩個機櫃編號只存在圖形裡)。給它
        # 們一個自己的小節,歸屬才不會被誤讀(2026-08-19 code review)
        if extras and split:
            blocks.append(Heading(2, f"{ws.title}:整張工作表的圖片與圖形文字"))
        blocks.extend(extras)
    # 巨集放最後:它是附屬內容,不該把工作表的資料擠到下面去
    blocks.extend(_macro_blocks(src))
    return blocks


def _macro_blocks(src: Path) -> list[Block]:
    """含巨集就給大綱,讀不出來就給標記。**docm / pptm / xlsm 共用**。

    `.docm`/`.pptm` 與 `.docx`/`.pptx` 在格式上完全相同,zip 裡只多一個
    `vbaProject.bin` ——所以共用同一個 reader 是正確的,而巨集這一段也
    只該有一份(`_has_macros` 本來就只是掃 zip 名稱,不綁 Excel)。"""
    if not _has_macros(src):
        return []
    return _vba_outline(src) or [Note(
        "本檔含巨集(VBA),但程式碼讀不出來"
        "——如果這份檔案的重點是那些程式,請直接開啟原始檔",
        docmd.KIND_MACRO,
    )]


def _sheet_images(ws, assets: AssetsDir, ocr_enabled: bool = True) -> list[Block]:
    """工作表裡的內嵌圖片(不是圖表),連同圖裡的文字。"""
    out: list[Block] = []
    for img in getattr(ws, "_images", None) or []:
        try:
            data = img._data()
        except Exception:
            logger.debug("Excel 內嵌圖取不到內容,略過", exc_info=True)
            continue
        link = assets.add_bytes(data, ".png")
        if link:
            out.append(Image(link, f"{ws.title} 的內嵌圖片"))
        out.extend(docimage.ocr_image_bytes(
            data, f"{ws.title} 的內嵌圖片", ocr_enabled,
        ))
    return out
